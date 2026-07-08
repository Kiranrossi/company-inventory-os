#!/usr/bin/env python3
import os
import sys

def main():
    try:
        import docx
    except ImportError:
        print("docx library not found. Please activate the virtual environment first.", file=sys.stderr)
        sys.exit(1)
        
    doc = docx.Document()
    doc.add_heading('Sample Purchase Order - Modulr Homes', 0)
    
    # Add a paragraph
    doc.add_paragraph('Please supply the following materials for Project Apollo 20:')
    
    # Add a table with wrapping test and numeric test
    table = doc.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Item No.'
    hdr_cells[1].text = 'Material Description'
    hdr_cells[2].text = 'Model / Spec'
    hdr_cells[3].text = 'Quantity'
    
    items = [
        ("1", "8mm CSMR - Century Sainik", "Century Sainik\nPremium Board", "12"),  # Exact / wrapped cell test
        ("2", "16mm CSMR - Century Sainik", "Sainik MR Ply\n16mm thickness", "5"),  # Fuzzy match test (score > 75)
        ("3", "2 x 22 mm 6952 Off White Edgeband", "Off White edge-band roll", "1500"),  # Edgeband length conversion (1500mm -> 1.5m)
        ("4", "Hettich Skirting Legs 100mm", "Skirting legs\nHettich brand", "8"),  # Fuzzy match test (score > 75)
        ("5", "Random Unmatching Core Material", "Non-existent item", "20"),  # Unmatched item test
        ("6", "Invalid Quantity Row", "Test skip row", "not_a_number"),  # Invalid quantity test (should be skipped)
    ]
    
    for item_no, mat, model, qty in items:
        row_cells = table.add_row().cells
        row_cells[0].text = item_no
        row_cells[1].text = mat
        row_cells[2].text = model
        row_cells[3].text = qty
        
    os.makedirs("tmp", exist_ok=True)
    doc.save("tmp/sample_order.docx")
    print("Created tmp/sample_order.docx successfully!")

if __name__ == "__main__":
    main()
